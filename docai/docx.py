"""Convert docx files to markdown."""

import re
from typing import IO

import docx
from docx.table import Table
from tabulate import tabulate

from docai.settings import Settings


def _style_separator(paragraph_style: str, separator: str) -> str:
    # TODO: Review formatting rules
    match paragraph_style:
        case (
            "List"
            | "List 2"
            | "List 3"
            | "List Bullet"
            | "List Bullet 2"
            | "List Bullet 3"
            | "List Continue"
            | "List Continue 2"
            | "List Continue 3"
            | "List Number"
            | "List Number 2"
            | "List Number 3"
            | "List Paragraph"
        ):
            separator = "\n" + "- "
        case "Heading 1" | "Title":
            separator = "\n\n" + "# "
        case "Subheading" | "Heading 2":
            separator = "\n\n" + "## "
        case "Heading 3":
            separator = "\n\n" + "### "
        case "Heading 4":
            separator = "\n\n" + "#### "
        case "Heading 5" | "Heading 6" | "Heading 7" | "Heading 8" | "Heading 9":
            separator = "\n\n" + "##### "
        case _:
            pass
    return separator


def convert_docx(file: IO[bytes], settings: Settings | None = None) -> str:
    """Convert docx into documents."""
    if settings is None:
        settings = Settings()
    markdown = ""
    document = docx.Document(file)
    for content in document.iter_inner_content():
        separator = "\n\n"
        # TODO: Test nested tables
        if isinstance(content, Table):
            # Normalise whitespace characters to not break github tables
            tabular_data = [
                [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells] for row in content.rows
            ]
            markdown += separator + tabulate(
                tabular_data,
                tablefmt=settings.tables.tablefmt,
                showindex=settings.tables.showindex,
                headers="firstrow",
            )
        else:  # Inner contents can only be a Table or Paragraph
            paragraph_style = content.style.name if content.style is not None else None
            if paragraph_style is not None:
                separator = _style_separator(paragraph_style, separator)
            markdown += separator + content.text.strip()
    markdown = markdown.strip()
    return markdown
